import logging
from nfty.njson import dc, dumps
import openpyxl


logger = logging.getLogger("AppLogger")


def json2docx_b64(meta):
    from docx import Document
    from io import BytesIO
    from base64 import b64encode

    meta = dc(meta)
    meta = dumps(meta, indent=2)
    d = Document()
    d.add_paragraph(meta)
    s = BytesIO()
    d.save(s)
    return b64encode(s.getvalue())


def json2docx_file(meta, filename):
    from docx import Document

    meta = dc(meta)
    meta = dumps(meta, indent=2)
    d = Document()
    d.add_paragraph(meta)
    if ".docx" not in filename:
        filename = filename + ".docx"
    return d.save(filename)


class excel:
    # wb = 'C:\Temp\securityemail.xlsx'
    # ws = '2018'
    def __init__(self, file):
        self.file = file
        self.wb = None

    def open(self, binary=False, data_only=False):
        if binary:
            from io import BytesIO

            try:
                self.wb = openpyxl.load_workbook(
                    filename=BytesIO(self.file.read()), data_only=data_only
                )
            except:
                self.wb = openpyxl.load_workbook(
                    filename=BytesIO(self.file), data_only=data_only
                )
        else:
            self.wb = openpyxl.load_workbook(filename=self.file, data_only=data_only)
        return self

    def ws(self, sheetname=False, header=True):
        if not sheetname:
            ws = self.wb[self.wb.sheetnames[0]]
        else:
            ws = self.wb[sheetname]
        xlist = []
        xa = xlist.append
        for row in ws.rows:
            xcell = []
            for cell in row:
                xcell.append(cell.value)
            xa(xcell)
        if header == False:
            return xlist[1:]
        return xlist

    def ws_hyperlinks(self, sheetname=False, header=True):
        if not sheetname:
            ws = self.wb[self.wb.sheetnames[0]]
        else:
            ws = self.wb[sheetname]
        xlist = []
        xa = xlist.append
        for row in ws.rows:
            xcell = []
            for cell in row:
                try:
                    xcell.append(cell.hyperlink.target)
                except:
                    xcell.append(cell.value)
            xa(xcell)
        if header == False:
            return xlist[1:]
        return xlist
